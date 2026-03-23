# services/soap/app.py
"""
SOAP Note Generator Service
Generates structured clinical notes from transcripts using LLM
Persists notes in Postgres for clinician review and FHIR writeback.
"""
import os
import json
import logging
import base64
import io
import time
from datetime import datetime
from typing import Optional, List, Dict, Any

import asyncpg
import httpx
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from prometheus_client import Counter, Histogram, generate_latest, CONTENT_TYPE_LATEST
from starlette.responses import Response

from llm_client import LlmClient
from soap_pipeline import (
    apply_field_update,
    generate_field_value,
    generate_structured_note,
    get_field_value,
    parse_soap_lines,
    resolve_section_from_path,
    summarize_value,
)
from template_engine import has_placeholders
from templates_builtin import SYSTEM_TEMPLATES
from template_store import TemplateStore

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("soap")

def safe_print(*args, **_kwargs):
    logger.debug("suppressed print", extra={"fields": len(args)})

print = safe_print


def log_safe(level: int, msg: str, request: Request | None = None, session_id: str | None = None, **kwargs):
    extra = {
        "correlationId": request.headers.get("x-correlation-id") if request else None,
        "sessionId": session_id,
    }
    for k, v in kwargs.items():
        if v is not None:
            extra[k] = v
    logger.log(level, msg, extra=extra)

app = FastAPI(title="SOAP Generator Service", version="1.1.0")

# CORS: configurable via env, default to localhost only
CORS_ALLOWED_ORIGINS = os.getenv("CORS_ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:5173").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[o.strip() for o in CORS_ALLOWED_ORIGINS],
    allow_methods=["GET", "POST", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization", "x-internal-secret", "x-correlation-id", "x-tenant-id"],
)

# Optional OTEL
try:
    from otel_setup import init_otel
    init_otel("soap", app=app)
except Exception:
    logger.debug("OTEL init skipped for SOAP")

soap_requests_total = Counter(
    "soap_requests_total",
    "Total SOAP service requests",
    ["endpoint", "status"],
)
soap_latency_seconds = Histogram(
    "soap_latency_seconds",
    "SOAP service request latency",
    ["endpoint", "status"],
    buckets=[0.05, 0.1, 0.25, 0.5, 1, 2, 3, 5, 10],
)

INTERNAL_SECRET = os.getenv("INTERNAL_SECRET", "")
if not INTERNAL_SECRET:
    raise RuntimeError("INTERNAL_SECRET must be set for SOAP service")
def _require_env(keys):
    missing = [k for k in keys if not os.getenv(k)]
    if missing:
        raise RuntimeError(f"Missing required env: {', '.join(missing)}")

_require_env(["DATABASE_URL"])

@app.middleware("http")
async def internal_auth(request: Request, call_next):
    if (
        request.url.path.startswith("/health")
        or request.url.path.startswith("/ready")
        or request.url.path.startswith("/metrics")
    ):
        return await call_next(request)
    # Use constant-time comparison to prevent timing attacks
    import hmac
    provided_secret = request.headers.get("x-internal-secret") or ""
    if not INTERNAL_SECRET or not hmac.compare_digest(provided_secret, INTERNAL_SECRET):
        raise HTTPException(status_code=401, detail="Unauthorized")
    return await call_next(request)

@app.middleware("http")
async def metrics_middleware(request: Request, call_next):
    if request.url.path.startswith("/metrics"):
        return Response(content=generate_latest(), media_type=CONTENT_TYPE_LATEST)
    start = time.time()
    response = await call_next(request)
    status = "ok" if response.status_code < 400 else "error"
    soap_requests_total.labels(endpoint=request.url.path, status=status).inc()
    soap_latency_seconds.labels(endpoint=request.url.path, status=status).observe(time.time() - start)
    return response

LLM_SERVICE_URL = os.getenv("LLM_SERVICE_URL", "http://localhost:5001")
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://postgres:postgres@postgres:5432/healthtech")
_pool: Optional[asyncpg.pool.Pool] = None
_llm_client: Optional[LlmClient] = None
_template_store: Optional[TemplateStore] = None

DOC_SUMMARY_SYSTEM_PROMPT = (
    "You are a clinical summarizer. Read the patient document text and provide a concise summary in English.\n"
    "Include: key conditions, medications, allergies, relevant dates, and recent procedures if present.\n"
    "Use 3-5 short sentences. Do not add new information."
)
MAX_DOCUMENT_CHARS = 20000

# ---------------------------
# Models
# ---------------------------
class SOAPRequest(BaseModel):
    transcript: str
    sessionId: Optional[str] = None
    patientContext: Optional[dict] = None
    patientId: str
    clinicianId: str
    encounterId: Optional[str] = None
    templateId: Optional[str] = None
    templateJson: Optional[Dict[str, Any]] = None
    patientName: Optional[str] = None
    providerName: Optional[str] = None
    dateOfVisit: Optional[str] = None

class SOAPResponse(BaseModel):
    id: str
    session_id: str
    patient_id: Optional[str]
    clinician_id: Optional[str]
    template_id: Optional[str] = None
    status: str
    subjective: str
    objective: str
    assessment: str
    plan: str
    icd_codes: Optional[List[str]] = None
    cpt_codes: Optional[List[str]] = None
    created_at: Optional[datetime] = None
    updated_at: Optional[datetime] = None
    raw_transcript: Optional[str] = None
    soap_json: Optional[Dict[str, Any]] = None

class ApproveRejectRequest(BaseModel):
    action: str  # approve|reject

class TemplateCreateRequest(BaseModel):
    name: str
    template: Dict[str, Any]
    id: Optional[str] = None

class TemplateResponse(BaseModel):
    id: str
    name: str
    template: Dict[str, Any]
    is_system: bool = False

class FieldUpdateRequest(BaseModel):
    fieldPath: str
    transcript: str
    mode: Optional[str] = "append"
    valueType: Optional[str] = None
    actorId: Optional[str] = None
    source: Optional[str] = None

class SectionsUpdateRequest(BaseModel):
    soapText: Optional[str] = None
    subjective: Optional[str] = None
    objective: Optional[str] = None
    assessment: Optional[str] = None
    plan: Optional[str] = None
    actorId: Optional[str] = None

class PatientCreateRequest(BaseModel):
    displayName: str
    externalId: Optional[str] = None

class PatientResponse(BaseModel):
    id: str
    display_name: Optional[str] = None
    external_id: Optional[str] = None
    created_at: Optional[datetime] = None

class PatientDocumentUploadRequest(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    contentBase64: Optional[str] = None
    fileName: Optional[str] = None
    contentType: Optional[str] = None
    source: Optional[str] = None
    summarize: Optional[bool] = True

class PatientDocumentResponse(BaseModel):
    id: str
    patient_id: str
    title: Optional[str] = None
    content_type: Optional[str] = None
    summary_text: Optional[str] = None
    created_at: Optional[datetime] = None

class PatientContextResponse(BaseModel):
    patient_id: str
    documents: List[Dict[str, Any]]
    recent_notes: List[Dict[str, Any]]
    rag_items: List[Dict[str, Any]]

# ---------------------------
# Startup / DB
# ---------------------------
@app.on_event("startup")
async def startup():
    global _pool, _llm_client, _template_store
    try:
        _pool = await asyncpg.create_pool(DATABASE_URL, min_size=1, max_size=5)
        async with _pool.acquire() as conn:
            await conn.execute("SELECT 1")
        logger.info("SOAP DB pool ready")
    except Exception as e:
        logger.warning("SOAP DB init failed", extra={"error": str(e)})
        _pool = None
    _llm_client = LlmClient(LLM_SERVICE_URL, INTERNAL_SECRET)
    _template_store = TemplateStore(_pool)
    try:
        await _template_store.ensure_system_templates(SYSTEM_TEMPLATES)
    except Exception as e:
        logger.warning("Template seed failed", extra={"error": str(e)})

@app.on_event("shutdown")
async def shutdown():
    global _pool
    if _pool:
        await _pool.close()
        _pool = None

# ---------------------------
# Routes
# ---------------------------
@app.get("/health")
async def health():
    return {"ok": True, "service": "soap-generator"}


@app.get("/ready")
async def ready():
    db_ok = False
    llm_ok = False
    if _pool:
        try:
            async with _pool.acquire() as conn:  # type: ignore
                await conn.execute("SELECT 1")
            db_ok = True
        except Exception as e:
            logger.warning("SOAP DB readiness check failed", extra={"error": str(e)})
    try:
        async with httpx.AsyncClient(timeout=3.0) as client:
            resp = await client.get(f"{LLM_SERVICE_URL}/health")
            llm_ok = resp.status_code == 200
    except Exception:
        llm_ok = False

    return {"ready": db_ok and llm_ok, "db": db_ok, "llm": llm_ok}

@app.get("/templates")
async def list_templates():
    if not _template_store:
        return {"templates": []}
    templates = await _template_store.list_templates()
    return {"templates": templates}

@app.get("/templates/{template_id}", response_model=TemplateResponse)
async def get_template(template_id: str):
    if not _template_store:
        raise HTTPException(status_code=503, detail="Template store unavailable")
    template = await _template_store.get_template(template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    return template

@app.post("/templates")
async def create_template(req: TemplateCreateRequest):
    if not _template_store:
        raise HTTPException(status_code=503, detail="Template store unavailable")
    validate_template(req.template)
    template_id = await _template_store.create_template(
        name=req.name,
        template=req.template,
        created_by=None,
        template_id=req.id,
    )
    return {"id": template_id}

@app.post("/generate", response_model=SOAPResponse)
async def generate_soap(req: SOAPRequest, request: Request):
    if _pool is None:
        raise HTTPException(status_code=503, detail="DB not available")
    if _llm_client is None:
        raise HTTPException(status_code=503, detail="LLM client not ready")
    if not req.transcript:
        raise HTTPException(status_code=400, detail="transcript required")

    # Capture tenant_id and clinician_id from gateway-forwarded headers
    tenant_id = request.headers.get("x-tenant-id", "default")
    actor_id  = request.headers.get("x-user-id") or req.clinicianId

    session_id = req.sessionId or f"soap-{int(datetime.utcnow().timestamp())}"
    template_id = req.templateId
    template = await resolve_template(template_id, req.templateJson)
    patient_name = req.patientName or "[Patient Name]"
    provider_name = req.providerName or "[Provider Name]"
    date_of_visit = req.dateOfVisit or ""
    patient_context = req.patientContext
    if not patient_context and req.patientId:
        patient_context = await build_patient_context(req.patientId)

    log_safe(
        logging.INFO,
        "SOAP generation request",
        session_id=session_id,
        patient_id=req.patientId,
        clinician_id=req.clinicianId,
    )

    try:
        result = await generate_structured_note(
            _llm_client,
            req.transcript,
            template,
            patient_name,
            date_of_visit,
            provider_name,
            session_id,
            patient_context,
        )
        sections = result["sections"]
        note_struct = result["note_json"]
        note_codes = result.get("codes") or {}
    except httpx.TimeoutException:
        log_safe(logging.ERROR, "SOAP generation timeout talking to LLM", session_id=session_id)
        raise HTTPException(status_code=504, detail="LLM service timeout")
    except Exception as e:
        log_safe(
            logging.ERROR,
            "SOAP generation failed",
            session_id=session_id,
            error=str(type(e).__name__),
        )
        raise HTTPException(status_code=500, detail="SOAP generation failed")

    record = await save_note({
        "session_id": session_id,
        "patient_id": req.patientId,
        "clinician_id": req.clinicianId,
        "template_id": template_id,
        "encounter_id": req.encounterId,
        "status": "pending",
        "raw_transcript": req.transcript,
        "soap_json": note_struct,
        "subjective": sections.subjective,
        "objective": sections.objective,
        "assessment": sections.assessment,
        "plan": sections.plan,
        "icd_codes": note_codes.get("icd_codes") or (note_struct.get("icd_codes") if isinstance(note_struct, dict) else []) or [],
        "cpt_codes": note_codes.get("cpt_codes") or (note_struct.get("cpt_codes") if isinstance(note_struct, dict) else []) or [],
        "tenant_id": tenant_id,
        "actor_id": actor_id,
    })
    await store_note_rag_items(record)

    return record

@app.get("/patients")
async def list_patients_route():
    return {"patients": await list_patients()}

@app.post("/patients", response_model=PatientResponse)
async def create_patient_route(req: PatientCreateRequest):
    if not req.displayName:
        raise HTTPException(status_code=400, detail="displayName required")
    return await create_patient(req.displayName.strip(), req.externalId)

@app.get("/patients/{patient_id}/documents")
async def list_patient_documents_route(patient_id: str):
    return {"documents": await fetch_patient_documents(patient_id, limit=50)}

@app.post("/patients/{patient_id}/documents", response_model=PatientDocumentResponse)
async def upload_patient_document(patient_id: str, req: PatientDocumentUploadRequest):
    if not patient_id:
        raise HTTPException(status_code=400, detail="patient_id required")
    content_text, resolved_content_type = extract_document_text(req)
    if not content_text:
        raise HTTPException(status_code=400, detail="content empty")
    title = req.title or title_from_filename(req.fileName)
    summary_text = None
    if req.summarize and _llm_client is not None:
        summary_text = await summarize_document(_llm_client, content_text, patient_id)
    doc = await insert_patient_document(
        patient_id=patient_id,
        title=title,
        content_text=content_text,
        content_type=resolved_content_type or req.contentType,
        source=req.source,
        summary_text=summary_text,
    )
    if summary_text:
        await insert_patient_rag_item(
            patient_id,
            "doc_summary",
            doc.title or "Document Summary",
            summary_text,
            metadata={"document_id": doc.id},
            source_id=doc.id,
        )
    else:
        await insert_patient_rag_item(
            patient_id,
            "document",
            doc.title or "Document",
            truncate_text(content_text, 800),
            metadata={"document_id": doc.id},
            source_id=doc.id,
        )
    return doc

@app.post("/patients/{patient_id}/documents/{doc_id}/summary", response_model=PatientDocumentResponse)
async def summarize_patient_document(patient_id: str, doc_id: str):
    if _llm_client is None:
        raise HTTPException(status_code=503, detail="LLM client not ready")
    doc = await fetch_patient_document_raw(doc_id)
    if not doc or doc.get("patient_id") != patient_id:
        raise HTTPException(status_code=404, detail="Document not found")
    summary_text = await summarize_document(_llm_client, doc.get("content_text", ""), doc_id)
    updated = await update_patient_document_summary(doc_id, summary_text)
    await insert_patient_rag_item(
        patient_id,
        "doc_summary",
        updated.title or "Document Summary",
        summary_text,
        metadata={"document_id": updated.id},
        source_id=updated.id,
    )
    return updated

@app.get("/patients/{patient_id}/context", response_model=PatientContextResponse)
async def get_patient_context(patient_id: str):
    context = await build_patient_context(patient_id)
    return PatientContextResponse(
        patient_id=patient_id,
        documents=context.get("documents", []),
        recent_notes=context.get("recent_notes", []),
        rag_items=context.get("rag_items", []),
    )

@app.get("/patients/{patient_id}/rag")
async def list_patient_rag(patient_id: str):
    return {"items": await list_patient_rag_items(patient_id, limit=50)}

@app.get("/notes")
async def list_notes(request: Request, status: Optional[str] = None, clinicianId: Optional[str] = None):
    if _pool is None:
        raise HTTPException(status_code=503, detail="DB not available")
    tenant_id = request.headers.get("x-tenant-id", "default")
    rows = await fetch_notes(tenant_id=tenant_id, status=status, clinician_id=clinicianId)
    return {"notes": rows}

@app.get("/notes/{note_id}", response_model=SOAPResponse)
async def get_note(note_id: str, request: Request):
    tenant_id = request.headers.get("x-tenant-id", "default")
    note = await fetch_note(note_id, tenant_id=tenant_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    return note

@app.patch("/notes/{note_id}/approve", response_model=SOAPResponse)
async def approve_note(note_id: str, request: Request):
    tenant_id = request.headers.get("x-tenant-id", "default")
    updated = await update_status(note_id, "approved", tenant_id=tenant_id)
    if not updated:
        raise HTTPException(status_code=404, detail="Note not found")
    return updated

@app.patch("/notes/{note_id}/reject", response_model=SOAPResponse)
async def reject_note(note_id: str, request: Request):
    tenant_id = request.headers.get("x-tenant-id", "default")
    updated = await update_status(note_id, "rejected", tenant_id=tenant_id)
    if not updated:
        raise HTTPException(status_code=404, detail="Note not found")
    return updated

@app.patch("/notes/{note_id}/field", response_model=SOAPResponse)
async def update_note_field(note_id: str, req: FieldUpdateRequest):
    if _pool is None:
        raise HTTPException(status_code=503, detail="DB not available")
    if _llm_client is None:
        raise HTTPException(status_code=503, detail="LLM client not ready")
    if not req.fieldPath or not req.transcript:
        raise HTTPException(status_code=400, detail="fieldPath and transcript required")

    note = await fetch_note(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    note_json = note.soap_json if isinstance(note.soap_json, dict) else {}
    try:
        existing = get_field_value(note_json, req.fieldPath)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid field path")
    expected_type = (req.valueType or ("list" if isinstance(existing, list) else "string")).lower()
    if expected_type not in ("string", "list"):
        expected_type = "string"
    mode = (req.mode or "append").lower()
    if mode not in ("append", "replace"):
        mode = "append"

    try:
        new_value = await generate_field_value(
            _llm_client,
            req.fieldPath,
            req.transcript,
            expected_type,
            existing,
            note.session_id,
        )
        apply_field_update(note_json, req.fieldPath, new_value, mode)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid field path")

    updated_subjective = note.subjective or ""
    updated_objective = note.objective or ""
    updated_assessment = note.assessment or ""
    updated_plan = note.plan or ""
    section = resolve_section_from_path(req.fieldPath)
    summary = summarize_value(new_value)
    if section == "subjective":
        updated_subjective = merge_section_text(updated_subjective, summary, mode)
    elif section == "objective":
        updated_objective = merge_section_text(updated_objective, summary, mode)
    elif section == "assessment":
        updated_assessment = merge_section_text(updated_assessment, summary, mode)
    elif section == "plan":
        updated_plan = merge_section_text(updated_plan, summary, mode)

    raw_transcript = (note.raw_transcript or "").strip()
    if req.transcript.strip():
        addendum = req.transcript.strip()
        raw_transcript = f"{raw_transcript}\n\n[Addendum]\n{addendum}" if raw_transcript else addendum

    updated = await update_note_field_content(
        note_id,
        note_json,
        updated_subjective,
        updated_objective,
        updated_assessment,
        updated_plan,
        raw_transcript,
    )
    await insert_audit(
        note_id,
        req.actorId or "unknown",
        "SOAP_NOTE_FIELD_UPDATED",
        {
            "fieldPath": req.fieldPath,
            "mode": mode,
            "valueType": expected_type,
            "source": req.source or "voice",
        },
    )
    await store_note_rag_items(updated)
    return updated

@app.patch("/notes/{note_id}/sections", response_model=SOAPResponse)
async def update_note_sections(note_id: str, req: SectionsUpdateRequest):
    if _pool is None:
        raise HTTPException(status_code=503, detail="DB not available")
    note = await fetch_note(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")

    sections = None
    if req.soapText:
        sections = parse_soap_lines(req.soapText)
    subjective = (req.subjective or (sections.subjective if sections else None) or note.subjective or "").strip()
    objective = (req.objective or (sections.objective if sections else None) or note.objective or "").strip()
    assessment = (req.assessment or (sections.assessment if sections else None) or note.assessment or "").strip()
    plan = (req.plan or (sections.plan if sections else None) or note.plan or "").strip()

    if not any([subjective, objective, assessment, plan]):
        raise HTTPException(status_code=400, detail="No updates provided")

    updated = await update_note_sections_content(note_id, subjective, objective, assessment, plan)
    await insert_audit(
        note_id,
        req.actorId or "unknown",
        "SOAP_NOTE_SECTIONS_UPDATED",
        {"source": "manual"},
    )
    await store_note_rag_items(updated)
    return updated

# ---------------------------
# Helpers
# ---------------------------
def validate_template(template: Dict[str, Any]):
    if not isinstance(template, (dict, list)):
        raise HTTPException(status_code=400, detail="template must be a JSON object or array")
    raw = json.dumps(template, ensure_ascii=False)
    if len(raw) > 100_000:
        raise HTTPException(status_code=400, detail="template too large")
    if not has_placeholders(template):
        logger.warning("Template has no placeholders; output will be mostly static")


async def resolve_template(template_id: Optional[str], template_json: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    if template_json:
        validate_template(template_json)
        return template_json
    if template_id and _template_store:
        template = await _template_store.get_template(template_id)
        if template:
            return template["template"]
    return SYSTEM_TEMPLATES[0]["template"]

def merge_section_text(existing: str, incoming: str, mode: str) -> str:
    existing = (existing or "").strip()
    incoming = (incoming or "").strip()
    if not incoming:
        return existing
    if mode == "replace" or not existing or existing.lower() == "not mentioned.":
        return incoming
    if incoming.lower() in existing.lower():
        return existing
    if existing.endswith((".", "!", "?")):
        return f"{existing} {incoming}"
    return f"{existing}; {incoming}"

def truncate_text(text: str, limit: int = 1200) -> str:
    cleaned = (text or "").strip()
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[:limit].rsplit(" ", 1)[0] + "..."

def title_from_filename(file_name: str | None) -> str | None:
    if not file_name:
        return None
    base = os.path.basename(file_name)
    stem, _ = os.path.splitext(base)
    return stem or None

def normalize_base64_payload(payload: str) -> str:
    cleaned = (payload or "").strip()
    if cleaned.startswith("data:") and "," in cleaned:
        cleaned = cleaned.split(",", 1)[1]
    return cleaned

def decode_base64_payload(payload: str) -> bytes:
    cleaned = normalize_base64_payload(payload)
    if not cleaned:
        return b""
    try:
        return base64.b64decode(cleaned, validate=True)
    except Exception:
        padded = cleaned + "=" * (-len(cleaned) % 4)
        return base64.b64decode(padded)

def decode_text_bytes(payload: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return payload.decode(encoding)
        except Exception:
            continue
    return payload.decode("utf-8", errors="ignore")

def guess_content_type(content_type: str | None, file_name: str | None) -> str | None:
    if content_type:
        return content_type
    if not file_name:
        return None
    ext = os.path.splitext(file_name)[1].lower()
    return {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".json": "application/json",
        ".csv": "text/csv",
    }.get(ext)

def detect_document_kind(content_type: str | None, file_name: str | None) -> str:
    ct = (content_type or "").lower()
    ext = os.path.splitext(file_name or "")[1].lower()
    if ct == "application/msword" or ext == ".doc":
        return "doc"
    if "pdf" in ct or ext == ".pdf":
        return "pdf"
    if "wordprocessingml" in ct or ext == ".docx":
        return "docx"
    if ct.startswith("text/") or ext in (".txt", ".json", ".csv"):
        return "text"
    return "text"

def extract_pdf_text(payload: bytes) -> str:
    if PdfReader is None:
        raise RuntimeError("pypdf not installed")
    reader = PdfReader(io.BytesIO(payload))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n".join(parts).strip()

def extract_docx_text(payload: bytes) -> str:
    if DocxDocument is None:
        raise RuntimeError("python-docx not installed")
    doc = DocxDocument(io.BytesIO(payload))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()

def extract_document_text(req: PatientDocumentUploadRequest) -> tuple[str, str | None]:
    content_text = (req.content or "").strip()
    file_name = os.path.basename(req.fileName) if req.fileName else None
    resolved_content_type = guess_content_type(req.contentType, file_name)

    if req.contentBase64:
        try:
            raw_bytes = decode_base64_payload(req.contentBase64)
        except Exception:
            raise HTTPException(status_code=400, detail="contentBase64 invalid")
        if raw_bytes:
            kind = detect_document_kind(resolved_content_type, file_name)
            try:
                if kind == "pdf":
                    extracted = extract_pdf_text(raw_bytes)
                    if extracted.strip():
                        content_text = extracted.strip()
                    resolved_content_type = resolved_content_type or "application/pdf"
                elif kind == "docx":
                    extracted = extract_docx_text(raw_bytes)
                    if extracted.strip():
                        content_text = extracted.strip()
                    resolved_content_type = (
                        resolved_content_type
                        or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                elif kind == "doc":
                    raise HTTPException(status_code=415, detail="DOC files not supported")
                else:
                    extracted = decode_text_bytes(raw_bytes)
                    if extracted.strip():
                        content_text = extracted.strip()
                    resolved_content_type = resolved_content_type or "text/plain"
            except HTTPException:
                raise
            except Exception as exc:
                logger.warning("Document extraction failed", extra={"error": str(exc)})

    if not content_text:
        return "", resolved_content_type
    return truncate_text(content_text, MAX_DOCUMENT_CHARS), resolved_content_type

async def summarize_document(llm: LlmClient, content: str, session_id: str | None) -> str:
    content = truncate_text(content, 4000)
    if not content:
        return ""
    messages = [
        {"role": "system", "content": DOC_SUMMARY_SYSTEM_PROMPT},
        {"role": "user", "content": f"Patient document:\n{content}\n\nSummary:"},
    ]
    try:
        summary = await llm.generate(messages, max_new_tokens=180, temperature=0.0, session_id=session_id)
        return truncate_text(summary, 800)
    except Exception:
        return truncate_text(content, 800)

def build_note_summary(note: SOAPResponse) -> str:
    parts = []
    if note.subjective:
        parts.append(f"Subjective: {note.subjective}")
    if note.objective:
        parts.append(f"Objective: {note.objective}")
    if note.assessment:
        parts.append(f"Assessment: {note.assessment}")
    if note.plan:
        parts.append(f"Plan: {note.plan}")
    return truncate_text(" | ".join(parts), 1200)

async def insert_patient_rag_item(
    patient_id: str,
    item_type: str,
    title: str,
    content_text: str,
    metadata: Dict[str, Any] | None = None,
    source_id: str | None = None,
) -> None:
    if _pool is None or not patient_id or not content_text:
        return
    async with _pool.acquire() as conn:  # type: ignore
        await conn.execute(
            """
            INSERT INTO patient_rag_items (patient_id, item_type, title, content_text, metadata, source_id)
            VALUES ($1,$2,$3,$4,$5,$6)
            ON CONFLICT (patient_id, item_type, source_id)
            DO UPDATE SET title=EXCLUDED.title, content_text=EXCLUDED.content_text, metadata=EXCLUDED.metadata, created_at=now()
            """,
            patient_id,
            item_type,
            title,
            content_text,
            json.dumps(metadata or {}),
            source_id,
        )

async def list_patient_rag_items(patient_id: str, limit: int = 8) -> List[Dict[str, Any]]:
    if _pool is None:
        return []
    async with _pool.acquire() as conn:  # type: ignore
        rows = await conn.fetch(
            """
            SELECT id, patient_id, item_type, title, content_text, metadata, source_id, created_at
            FROM patient_rag_items
            WHERE patient_id=$1
            ORDER BY created_at DESC
            LIMIT $2
            """,
            patient_id,
            limit,
        )
    items = []
    for row in rows:
        items.append(
            {
                "id": str(row["id"]),
                "patient_id": row["patient_id"],
                "item_type": row["item_type"],
                "title": row["title"],
                "content_text": row["content_text"],
                "metadata": row.get("metadata") or {},
                "source_id": row.get("source_id"),
                "created_at": row.get("created_at"),
            }
        )
    return items

async def store_note_rag_items(note: SOAPResponse) -> None:
    if not note.patient_id:
        return
    summary = build_note_summary(note)
    await insert_patient_rag_item(
        note.patient_id,
        "soap_note",
        "SOAP Note",
        summary,
        metadata={"note_id": note.id, "created_at": str(note.created_at)},
        source_id=note.id,
    )
    if note.raw_transcript:
        await insert_patient_rag_item(
            note.patient_id,
            "transcript",
            "Transcript",
            truncate_text(note.raw_transcript, 1200),
            metadata={"note_id": note.id},
            source_id=note.id,
        )

async def list_patients() -> List[Dict[str, Any]]:
    if _pool is None:
        return []
    async with _pool.acquire() as conn:  # type: ignore
        rows = await conn.fetch(
            "SELECT id, display_name, external_id, created_at FROM patients ORDER BY created_at DESC"
        )
    return [
        {
            "id": str(row["id"]),
            "display_name": row.get("display_name"),
            "external_id": row.get("external_id"),
            "created_at": row.get("created_at"),
        }
        for row in rows
    ]

async def create_patient(display_name: str, external_id: str | None) -> PatientResponse:
    if _pool is None:
        raise HTTPException(status_code=503, detail="DB not available")
    async with _pool.acquire() as conn:  # type: ignore
        row = await conn.fetchrow(
            """
            INSERT INTO patients (display_name, external_id)
            VALUES ($1,$2)
            RETURNING id, display_name, external_id, created_at
            """,
            display_name,
            external_id,
        )
    return PatientResponse(
        id=str(row["id"]),
        display_name=row.get("display_name"),
        external_id=row.get("external_id"),
        created_at=row.get("created_at"),
    )

async def fetch_patient_documents(patient_id: str, limit: int = 10) -> List[Dict[str, Any]]:
    if _pool is None:
        return []
    async with _pool.acquire() as conn:  # type: ignore
        rows = await conn.fetch(
            """
            SELECT id, patient_id, title, content_type, summary_text, created_at
            FROM patient_documents
            WHERE patient_id=$1
            ORDER BY created_at DESC
            LIMIT $2
            """,
            patient_id,
            limit,
        )
    docs = []
    for row in rows:
        docs.append(
            {
                "id": str(row["id"]),
                "patient_id": row["patient_id"],
                "title": row.get("title"),
                "content_type": row.get("content_type"),
                "summary_text": row.get("summary_text"),
                "created_at": row.get("created_at"),
            }
        )
    return docs

async def fetch_patient_document_raw(doc_id: str) -> Optional[Dict[str, Any]]:
    if _pool is None:
        return None
    async with _pool.acquire() as conn:  # type: ignore
        row = await conn.fetchrow(
            """
            SELECT id, patient_id, title, content_text, content_type, summary_text, created_at
            FROM patient_documents
            WHERE id=$1
            """,
            doc_id,
        )
    if not row:
        return None
    return {
        "id": str(row["id"]),
        "patient_id": row["patient_id"],
        "title": row.get("title"),
        "content_text": row.get("content_text"),
        "content_type": row.get("content_type"),
        "summary_text": row.get("summary_text"),
        "created_at": row.get("created_at"),
    }

async def insert_patient_document(
    patient_id: str,
    title: str | None,
    content_text: str,
    content_type: str | None,
    source: str | None,
    summary_text: str | None,
) -> PatientDocumentResponse:
    if _pool is None:
        raise HTTPException(status_code=503, detail="DB not available")
    async with _pool.acquire() as conn:  # type: ignore
        row = await conn.fetchrow(
            """
            INSERT INTO patient_documents (patient_id, title, content_text, content_type, source, summary_text)
            VALUES ($1,$2,$3,$4,$5,$6)
            RETURNING id, patient_id, title, content_type, summary_text, created_at
            """,
            patient_id,
            title,
            content_text,
            content_type,
            source,
            summary_text,
        )
    return PatientDocumentResponse(
        id=str(row["id"]),
        patient_id=row["patient_id"],
        title=row.get("title"),
        content_type=row.get("content_type"),
        summary_text=row.get("summary_text"),
        created_at=row.get("created_at"),
    )

async def update_patient_document_summary(doc_id: str, summary_text: str) -> PatientDocumentResponse:
    if _pool is None:
        raise HTTPException(status_code=503, detail="DB not available")
    async with _pool.acquire() as conn:  # type: ignore
        row = await conn.fetchrow(
            """
            UPDATE patient_documents
            SET summary_text=$2
            WHERE id=$1
            RETURNING id, patient_id, title, content_type, summary_text, created_at
            """,
            doc_id,
            summary_text,
        )
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    return PatientDocumentResponse(
        id=str(row["id"]),
        patient_id=row["patient_id"],
        title=row.get("title"),
        content_type=row.get("content_type"),
        summary_text=row.get("summary_text"),
        created_at=row.get("created_at"),
    )

async def fetch_recent_notes_by_patient(patient_id: str, limit: int = 3) -> List[Dict[str, Any]]:
    if _pool is None:
        return []
    async with _pool.acquire() as conn:  # type: ignore
        rows = await conn.fetch(
            """
            SELECT id, subjective, objective, assessment, plan, created_at, raw_transcript
            FROM soap_notes
            WHERE patient_id=$1
            ORDER BY created_at DESC
            LIMIT $2
            """,
            patient_id,
            limit,
        )
    notes = []
    for row in rows:
        summary = truncate_text(
            " | ".join(
                [
                    f"Subjective: {row.get('subjective') or ''}",
                    f"Assessment: {row.get('assessment') or ''}",
                    f"Plan: {row.get('plan') or ''}",
                ]
            ),
            800,
        )
        notes.append(
            {
                "id": str(row["id"]),
                "summary": summary,
                "created_at": row.get("created_at"),
                "transcript": truncate_text(row.get("raw_transcript") or "", 500),
            }
        )
    return notes

async def build_patient_context(patient_id: str) -> Dict[str, Any]:
    if not patient_id:
        return {}
    documents = await fetch_patient_documents(patient_id, limit=3)
    recent_notes = await fetch_recent_notes_by_patient(patient_id, limit=3)
    rag_items = await list_patient_rag_items(patient_id, limit=6)
    return {
        "patient_id": patient_id,
        "documents": documents,
        "recent_notes": recent_notes,
        "rag_items": rag_items,
    }

async def insert_audit(note_id: str, actor_id: str, action: str, metadata: Dict[str, Any]):
    if _pool is None:
        return
    async with _pool.acquire() as conn:  # type: ignore
        await conn.execute(
            "INSERT INTO soap_note_audit (note_id, actor_id, action, metadata) VALUES ($1,$2,$3,$4)",
            note_id,
            actor_id,
            action,
            json.dumps(metadata or {}),
        )

async def update_note_field_content(
    note_id: str,
    soap_json: Dict[str, Any],
    subjective: str,
    objective: str,
    assessment: str,
    plan: str,
    raw_transcript: str,
) -> SOAPResponse:
    async with _pool.acquire() as conn:  # type: ignore
        row = await conn.fetchrow(
            """
            UPDATE soap_notes
            SET soap_json=$2, subjective=$3, objective=$4, assessment=$5, plan=$6, raw_transcript=$7, updated_at=now()
            WHERE id=$1
            RETURNING *
            """,
            note_id,
            json.dumps(soap_json),
            subjective,
            objective,
            assessment,
            plan,
            raw_transcript,
        )
    if not row:
        raise HTTPException(status_code=404, detail="Note not found")
    return record_to_model(row)

async def update_note_sections_content(
    note_id: str,
    subjective: str,
    objective: str,
    assessment: str,
    plan: str,
) -> SOAPResponse:
    async with _pool.acquire() as conn:  # type: ignore
        row = await conn.fetchrow(
            """
            UPDATE soap_notes
            SET subjective=$2, objective=$3, assessment=$4, plan=$5, updated_at=now()
            WHERE id=$1
            RETURNING *
            """,
            note_id,
            subjective,
            objective,
            assessment,
            plan,
        )
    if not row:
        raise HTTPException(status_code=404, detail="Note not found")
    return record_to_model(row)

async def save_note(note: Dict[str, Any]) -> SOAPResponse:
    if not note.get("patient_id") or not note.get("clinician_id") or not note.get("session_id"):
        raise HTTPException(status_code=400, detail="patient_id, clinician_id, session_id required")
    icd = note.get("icd_codes") or []
    cpt = note.get("cpt_codes") or []
    soap_json = note.get("soap_json") or {}
    tenant_id = note.get("tenant_id") or "default"
    actor_id  = note.get("actor_id")  or note.get("clinician_id") or "unknown"
    async with _pool.acquire() as conn:  # type: ignore
        # Atomically create the SOAP note AND the audit trail entry.
        # If either INSERT fails the whole transaction rolls back — no unaudited PHI.
        async with conn.transaction():
            # RLS: tell Postgres which tenant this request belongs to
            # asyncpg does not support parameters in SET LOCAL, so sanitize manually
            _safe_tid = "".join(c for c in (tenant_id or "default") if c.isalnum() or c in "-_.") or "default"
            await conn.execute(f"SET LOCAL app.tenant_id = '{_safe_tid}'")
            row = await conn.fetchrow(
                """
                INSERT INTO soap_notes (session_id, patient_id, clinician_id, template_id, status, raw_transcript, soap_json,
                                        subjective, objective, assessment, plan, icd_codes, cpt_codes)
                VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11,$12,$13)
                RETURNING id, session_id, patient_id, clinician_id, template_id, status, subjective, objective, assessment, plan, icd_codes, cpt_codes, created_at, updated_at, raw_transcript, soap_json
                """,
                note.get("session_id"),
                note.get("patient_id"),
                note.get("clinician_id"),
                note.get("template_id"),
                note.get("status", "pending"),
                note.get("raw_transcript"),
                json.dumps(soap_json),
                note.get("subjective"),
                note.get("objective"),
                note.get("assessment"),
                note.get("plan"),
                icd,
                cpt,
            )
            # Audit log insert — inside the same transaction
            await conn.execute(
                """
                INSERT INTO audit_log (actor_id, action, resource_type, resource_id, metadata, tenant_id)
                VALUES ($1, $2, $3, $4, $5::jsonb, $6)
                """,
                actor_id,
                "SOAP_NOTE_CREATED",
                "soap_note",
                str(row["id"]),
                json.dumps({
                    "session_id": note.get("session_id"),
                    "patient_id":  note.get("patient_id"),
                }),
                tenant_id,
            )
    return record_to_model(row)

async def fetch_notes(tenant_id: str = "default", status: Optional[str] = None, clinician_id: Optional[str] = None) -> List[Dict[str, Any]]:
    query = "SELECT * FROM soap_notes"
    conds: List[str] = []
    params: List[Any] = []
    if status:
        params.append(status)
        conds.append(f"status = ${len(params)}")
    if clinician_id:
        params.append(clinician_id)
        conds.append(f"clinician_id = ${len(params)}")
    if conds:
        query += " WHERE " + " AND ".join(conds)
    query += " ORDER BY created_at DESC"
    async with _pool.acquire() as conn:  # type: ignore
        async with conn.transaction():
            _safe_tid = "".join(c for c in (tenant_id or "default") if c.isalnum() or c in "-_.") or "default"
            await conn.execute(f"SET LOCAL app.tenant_id = '{_safe_tid}'")
            rows = await conn.fetch(query, *params)
    return [record_to_model(r).dict() for r in rows]

async def fetch_note(note_id: str, tenant_id: str = "default") -> Optional[SOAPResponse]:
    async with _pool.acquire() as conn:  # type: ignore
        async with conn.transaction():
            _safe_tid = "".join(c for c in (tenant_id or "default") if c.isalnum() or c in "-_.") or "default"
            await conn.execute(f"SET LOCAL app.tenant_id = '{_safe_tid}'")
            row = await conn.fetchrow("SELECT * FROM soap_notes WHERE id = $1", note_id)
    if not row:
        return None
    return record_to_model(row)

async def update_status(note_id: str, status: str, tenant_id: str = "default") -> Optional[SOAPResponse]:
    async with _pool.acquire() as conn:  # type: ignore
        async with conn.transaction():
            _safe_tid = "".join(c for c in (tenant_id or "default") if c.isalnum() or c in "-_.") or "default"
            await conn.execute(f"SET LOCAL app.tenant_id = '{_safe_tid}'")
            row = await conn.fetchrow(
                "UPDATE soap_notes SET status=$2, updated_at=now() WHERE id=$1 RETURNING *",
                note_id,
                status,
            )
    if not row:
        return None
    return record_to_model(row)


def record_to_model(row: asyncpg.Record) -> SOAPResponse:
    soap_json = row.get("soap_json")
    if isinstance(soap_json, str):
        try:
            soap_json = json.loads(soap_json)
        except Exception:
            soap_json = {}
    return SOAPResponse(
        id=str(row["id"]),
        session_id=row.get("session_id"),
        patient_id=row.get("patient_id"),
        clinician_id=row.get("clinician_id"),
        template_id=row.get("template_id"),
        status=row.get("status"),
        subjective=row.get("subjective") or "",
        objective=row.get("objective") or "",
        assessment=row.get("assessment") or "",
        plan=row.get("plan") or "",
        icd_codes=row.get("icd_codes"),
        cpt_codes=row.get("cpt_codes"),
        created_at=row.get("created_at"),
        updated_at=row.get("updated_at"),
        raw_transcript=row.get("raw_transcript"),
        soap_json=soap_json,
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5003, log_level="info")
